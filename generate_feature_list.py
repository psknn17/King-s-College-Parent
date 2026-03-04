#!/usr/bin/env python3
"""Generate Feature Lists .docx for King's College Parent Portal."""

import glob
import os
from docx import Document
from docx.shared import Pt, Twips, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Find output directory
dirs = glob.glob("/Users/passkornnabangchang/Desktop/King*College*")
if not dirs:
    raise RuntimeError("Could not find King's College directory on Desktop")
output_dir = dirs[0]
output_path = os.path.join(output_dir, "Feature lists - Parent Portal_04032026.docx")

doc = Document()

# ── Helper functions ──────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, size=14, bold=False, font_name="TH SarabunPSK"):
    """Configure a run's font properties."""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    # Set East-Asian and complex-script font
    r = run._element
    rPr = r.get_or_add_rPr()
    # Complex script font (for Thai)
    cs_font = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="{font_name}"/>')
    existing = rPr.findall(qn('w:rFonts'))
    if existing:
        existing[0].set(qn('w:cs'), font_name)
        existing[0].set(qn('w:eastAsia'), font_name)
    else:
        rPr.append(cs_font)
    if bold:
        # Ensure bold for complex script
        bCs = parse_xml(f'<w:bCs {nsdecls("w")} w:val="1"/>')
        rPr.append(bCs)

def add_styled_paragraph(doc_or_cell, text, size=14, bold=False, alignment=None):
    """Add a paragraph with styled text."""
    p = doc_or_cell.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    # Set paragraph-level font
    pPr = p._element.get_or_add_pPr()
    return p

def write_cell_text(cell, text, bold=False, size=14):
    """Write text into a cell, handling \\n as line breaks and preserving '- ' bullets."""
    # Clear existing paragraphs
    for p in cell.paragraphs:
        p.clear()

    # Use the first existing paragraph
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if not text:
        run = p.add_run("")
        set_run_font(run, size=size, bold=bold)
        return

    lines = text.split("\\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold)
        if i < len(lines) - 1:
            run.add_break()

def set_col_width(col_cells, width_twips):
    """Set width for column cells."""
    for cell in col_cells:
        cell.width = Twips(width_twips)

def set_cell_vertical_alignment(cell, val="top"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val}"/>')
    tcPr.append(vAlign)

def set_cell_margins(cell, top=0, bottom=0, left=57, right=57):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


# ── Set default document font ────────────────────────────────────

style = doc.styles['Normal']
font = style.font
font.name = "TH SarabunPSK"
font.size = Pt(14)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="TH SarabunPSK" w:hAnsi="TH SarabunPSK" w:cs="TH SarabunPSK" w:eastAsia="TH SarabunPSK"/>')
    rPr.append(rFonts)
else:
    rFonts.set(qn('w:ascii'), "TH SarabunPSK")
    rFonts.set(qn('w:hAnsi'), "TH SarabunPSK")
    rFonts.set(qn('w:cs'), "TH SarabunPSK")
    rFonts.set(qn('w:eastAsia'), "TH SarabunPSK")

# ── COVER PAGE ────────────────────────────────────────────────────

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("Feature lists")
set_run_font(run, size=16, bold=True)

# Subtitle
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_sub.add_run("Parent Portal")
set_run_font(run, size=14, bold=False)

# Blank lines
doc.add_paragraph()
doc.add_paragraph()

# Cover table
cover_table = doc.add_table(rows=3, cols=2)
cover_table.style = 'Table Grid'

# Row 1: Link
cell = cover_table.cell(0, 0)
write_cell_text(cell, "Link", bold=True)
cell = cover_table.cell(0, 1)
write_cell_text(cell, "")

# Row 2: Role
cell = cover_table.cell(1, 0)
write_cell_text(cell, "Role", bold=True)
cell = cover_table.cell(1, 1)
write_cell_text(cell, "1. Parent")

# Row 3: Module
cell = cover_table.cell(2, 0)
write_cell_text(cell, "Module", bold=True)
cell = cover_table.cell(2, 1)
write_cell_text(cell, "A. Login, B. Dashboard, C. Tuition, D. ECA, E. Trip & Activity, F. Exam, G. School Bus, H. Cart & Payment, I. Transaction History, J. Account Settings, K. Multi-Language, L. Responsive Design")

# Page break before main content
doc.add_page_break()

# ── MAIN TABLE DATA ───────────────────────────────────────────────

# Row types: 'module' = grey shading + bold, 'sub' = bold WBS+Feature, 'detail' = normal
# Format: (WBS, Feature, Description, row_type)
rows_data = [
    ("A", "เข้าสู่ระบบ (Login)", "", "module"),
    ("A1", "เข้าสู่ระบบ (SSO)", "", "sub"),
    ("A1.1", "เข้าสู่ระบบผ่านระบบโรงเรียน (SSO)", "ผู้ปกครองถูกยืนยันตัวตนจากระบบโรงเรียนมาแล้ว\\n- กดปุ่ม \"เข้าสู่ระบบ\" เข้า Portal ได้เลย\\n- ไม่ต้องกรอก Email / Password / OTP", "detail"),
    ("A1.2", "หน้า Welcome", "แสดงโลโก้โรงเรียน + ภาพพื้นหลัง\\n- ปุ่มเข้าสู่ระบบปุ่มเดียว\\n- ตัวเลือกภาษามุมขวาบน", "detail"),

    ("B", "แดชบอร์ด (Dashboard)", "", "module"),
    ("B1", "ภาพรวมข้อมูล (Summary Overview)", "", "sub"),
    ("B1.1", "แสดงยอดใบแจ้งหนี้ค้างชำระ", "Summary Box: จำนวนใบแจ้งหนี้ค้างชำระ + ยอดรวม\\nสีตามประเภท: primary, warning, destructive", "detail"),
    ("B1.2", "แสดงยอดชำระแล้วในเทอมนี้", "Summary Box: แสดงยอดเงินที่ชำระแล้วในเทอมปัจจุบัน", "detail"),
    ("B1.3", "แสดงยอด Credit Balance", "Summary Box: แสดงยอดเครดิตคงเหลือที่ใช้หักจากการชำระได้", "detail"),
    ("B1.4", "แสดงจำนวนกิจกรรมที่เปิดรับ", "Summary Box: แสดงจำนวนกิจกรรม/คอร์สที่เปิดลงทะเบียน", "detail"),
    ("B2", "ทางลัด (Quick Actions)", "", "sub"),
    ("B2.1", "ปุ่มทางลัดเข้าถึงฟีเจอร์หลัก", "4 ปุ่ม: Pay Now, Register ECA, View Events, View Receipts\\n- มีไอคอนและคำอธิบาย\\n- สีแยกตามประเภท", "detail"),
    ("B3", "การแจ้งเตือนและกำหนดชำระ (Notifications & Deadlines)", "", "sub"),
    ("B3.1", "แสดง Upcoming Deadlines", "Toggle เปิด/ปิดรายการกำหนดชำระ\\n- แสดงไอคอนตามประเภท (Tuition, ECA, Trip, Exam, School Bus)\\n- แสดงยอดเงินและวันครบกำหนด", "detail"),
    ("B3.2", "Popup แจ้งเตือนค้างชำระ (Overdue)", "Dialog แจ้งเตือนเมื่อมีใบแจ้งหนี้เลยกำหนดชำระ\\n- แสดงรายละเอียดยอดค้างชำระ\\n- ปุ่ม Pay Now", "detail"),
    ("B3.3", "Children Overview พร้อม Badge แจ้งเตือน", "แสดงภาพรวมบุตรทั้งหมด\\n- Badge แสดงจำนวนรายการค้างชำระ\\n- กดเพื่อดูรายละเอียดแต่ละคน", "detail"),

    ("C", "ค่าเทอม (Tuition)", "", "module"),
    ("C1", "ใบแจ้งหนี้ค่าเทอม (Tuition Invoices)", "", "sub"),
    ("C1.1", "แสดง Invoice จาก Backoffice", "แสดงรายการใบแจ้งหนี้ที่เจ้าหน้าที่สร้างจากระบบ Backoffice\\n- แยกกลุ่ม: ค้างชำระ / ชำระแล้ว\\n- Accordion เปิด/ปิดได้", "detail"),
    ("C1.2", "แสดงสถานะใบแจ้งหนี้", "สถานะ: Pending, Overdue, Paid, Partial\\n- Badge สีตามสถานะ\\n- แสดงวันครบกำหนดชำระ", "detail"),
    ("C1.3", "หยิบใบแจ้งหนี้ใส่ตะกร้า", "กดปุ่ม Add to Cart → เพิ่มเข้าตะกร้าเพื่อชำระเงิน", "detail"),
    ("C1.4", "กรองตามนักเรียน", "เลือกดูใบแจ้งหนี้ของนักเรียนแต่ละคน หรือดูทั้งหมด", "detail"),

    ("D", "กิจกรรมนอกเวลา (ECA)", "", "module"),
    ("D1", "ใบแจ้งหนี้ ECA (ECA Invoices)", "", "sub"),
    ("D1.1", "แสดง Invoice จาก Backoffice", "แสดงรายการใบแจ้งหนี้ ECA ที่เจ้าหน้าที่สร้างจากระบบ Backoffice\\n- แยกกลุ่ม: ค้างชำระ / ชำระแล้ว", "detail"),
    ("D1.2", "แสดงสถานะใบแจ้งหนี้", "สถานะ: Pending, Overdue, Paid\\n- Badge สีตามสถานะ", "detail"),
    ("D1.3", "หยิบใบแจ้งหนี้ใส่ตะกร้า", "กดปุ่ม Add to Cart → เพิ่มเข้าตะกร้า", "detail"),
    ("D1.4", "กรองตามนักเรียน", "เลือกดูใบแจ้งหนี้ของนักเรียนแต่ละคน", "detail"),

    ("E", "ทริปและกิจกรรม (Trip & Activity)", "", "module"),
    ("E1", "ใบแจ้งหนี้ทริป (Trip Invoices)", "", "sub"),
    ("E1.1", "แสดง Invoice จาก Backoffice", "แสดงรายการใบแจ้งหนี้ทริปที่เจ้าหน้าที่สร้างจากระบบ Backoffice\\n- แยกกลุ่ม: ค้างชำระ / ชำระแล้ว", "detail"),
    ("E1.2", "แสดงสถานะใบแจ้งหนี้", "สถานะ: Pending, Overdue, Paid\\n- Badge สีตามสถานะ", "detail"),
    ("E1.3", "หยิบใบแจ้งหนี้ใส่ตะกร้า", "กดปุ่ม Add to Cart → เพิ่มเข้าตะกร้า", "detail"),
    ("E1.4", "กรองตามนักเรียน", "เลือกดูใบแจ้งหนี้ของนักเรียนแต่ละคน", "detail"),

    ("F", "ค่าสอบ (Exam)", "", "module"),
    ("F1", "ใบแจ้งหนี้ค่าสอบ (Exam Invoices)", "", "sub"),
    ("F1.1", "แสดง Invoice จาก Backoffice", "แสดงรายการใบแจ้งหนี้ค่าสอบที่เจ้าหน้าที่สร้างจากระบบ Backoffice\\n- แยกกลุ่ม: ค้างชำระ / ชำระแล้ว", "detail"),
    ("F1.2", "แสดงสถานะใบแจ้งหนี้", "สถานะ: Pending, Overdue, Paid\\n- Badge สีตามสถานะ", "detail"),
    ("F1.3", "หยิบใบแจ้งหนี้ใส่ตะกร้า", "กดปุ่ม Add to Cart → เพิ่มเข้าตะกร้า", "detail"),
    ("F1.4", "กรองตามนักเรียน", "เลือกดูใบแจ้งหนี้ของนักเรียนแต่ละคน", "detail"),

    ("G", "รถรับส่ง (School Bus)", "", "module"),
    ("G1", "ใบแจ้งหนี้รถรับส่ง (School Bus Invoices)", "", "sub"),
    ("G1.1", "แสดง Invoice จาก Backoffice", "แสดงรายการใบแจ้งหนี้รถรับส่งที่เจ้าหน้าที่สร้างจากระบบ Backoffice\\n- แยกกลุ่ม: ค้างชำระ / ชำระแล้ว", "detail"),
    ("G1.2", "แสดงสถานะใบแจ้งหนี้", "สถานะ: Pending, Overdue, Paid\\n- Badge สีตามสถานะ", "detail"),
    ("G1.3", "หยิบใบแจ้งหนี้ใส่ตะกร้า", "กดปุ่ม Add to Cart → เพิ่มเข้าตะกร้า", "detail"),
    ("G1.4", "กรองตามนักเรียน", "เลือกดูใบแจ้งหนี้ของนักเรียนแต่ละคน", "detail"),

    ("H", "ตะกร้าสินค้าและชำระเงิน (Cart & Payment)", "", "module"),
    ("H1", "ตะกร้าสินค้า (Cart)", "", "sub"),
    ("H1.1", "แสดงรายการ Invoice ที่เลือกแยกตามนักเรียน", "จัดกลุ่มรายการตามชื่อนักเรียน\\n- แสดงชื่อ, ราคา", "detail"),
    ("H1.2", "เลือก/ยกเลิกรายการ", "Checkbox เลือก/ยกเลิกรายการแต่ละรายการ\\n- ปุ่ม Select All / Deselect All", "detail"),
    ("H1.3", "ลบรายการออกจากตะกร้า", "ปุ่ม X ลบรายการ\\n- อัปเดตยอดรวมทันที", "detail"),
    ("H1.4", "ใช้ Credit Note หักยอดชำระ", "Toggle เปิดใช้ Credit Note\\n- เลือก Credit Note ที่ต้องการใช้ได้หลายรายการ\\n- แสดงยอดเครดิตคงเหลือ\\n- คำนวณยอดหลังหักเครดิต", "detail"),
    ("H1.5", "คำนวณยอดรวม", "แสดง Subtotal และ Total\\n- อัปเดตอัตโนมัติเมื่อเปลี่ยนรายการ", "detail"),
    ("H2", "ชำระเงิน (Checkout)", "", "sub"),
    ("H2.1", "เลือกวิธีชำระเงิน", "รองรับ 2 วิธี:\\n1. Credit Card (Bank fee +2.9%)\\n2. Thai QR / PromptPay (Bank fee +0%, ฟรี)", "detail"),
    ("H2.2", "คำนวณค่าธรรมเนียมอัตโนมัติ", "แสดง Bank fee ตามวิธีชำระเงินที่เลือก\\n- คำนวณยอดรวมสุทธิ", "detail"),
    ("H2.3", "แสดง Payment Summary", "สรุปรายการ: ชื่อ, ราคา\\n- Subtotal, Credit Applied, Bank fee\\n- Total ยอดสุทธิ", "detail"),
    ("H2.4", "ใช้ Credit Note ในหน้า Checkout", "หักยอด Credit Note จากยอดชำระ\\n- แสดงรายละเอียด Credit ที่ใช้", "detail"),
    ("H3", "ชำระเงินสำเร็จ (Payment Success)", "", "sub"),
    ("H3.1", "แสดงหน้ายืนยันสำเร็จ", "ไอคอน Success\\n- เลขที่ใบเสร็จ\\n- ยอดเงิน, วันที่ชำระ\\n- วิธีชำระเงิน", "detail"),
    ("H3.2", "ดาวน์โหลดใบเสร็จ", "ปุ่ม Download Receipt → ดาวน์โหลดไฟล์ PDF", "detail"),
    ("H3.3", "กลับสู่หน้าหลัก", "ปุ่ม Back to Dashboard", "detail"),

    ("I", "ประวัติธุรกรรม (Transaction History)", "", "module"),
    ("I1", "ใบเสร็จรับเงิน (Receipts)", "", "sub"),
    ("I1.1", "แสดงรายการใบเสร็จ", "แสดง: วันที่, รายละเอียด, ยอดเงิน, วิธีชำระ, สถานะ\\n- Badge วิธีชำระ: Credit Card, Thai QR", "detail"),
    ("I1.2", "กรองใบเสร็จ", "กรองตาม: - นักเรียน\\n- ประเภท (Tuition, ECA, Trip & Activity, Exam, School Bus)\\n- เดือน", "detail"),
    ("I1.3", "ค้นหาใบเสร็จ", "ค้นหาจาก: คำอธิบาย, เลขที่อ้างอิง, Invoice ID", "detail"),
    ("I1.4", "ดาวน์โหลดใบเสร็จ PDF", "กดปุ่ม Download → สร้างและดาวน์โหลด PDF ใบเสร็จ\\n- รวมข้อมูลโรงเรียน, รายการชำระ, ข้อมูลผู้ชำระ\\n- จำนวนเงินเป็นตัวอักษร (Number to Words)", "detail"),
    ("I1.5", "แบ่งหน้าแสดงผล (Pagination)", "แสดง 5 รายการต่อหน้า\\n- ปุ่มเปลี่ยนหน้า", "detail"),
    ("I2", "ใบลดหนี้ (Credit Notes)", "", "sub"),
    ("I2.1", "แสดงรายการ Credit Note", "แสดง: เลขที่, ประเภท, ยอดเงิน, สถานะ\\n- ประเภท: Refund, Discount, Overpayment, Cancellation", "detail"),
    ("I2.2", "กรอง Credit Note", "กรองตาม:\\n- ปีการศึกษา\\n- ประเภท (Refund/Discount/Overpayment/Cancellation)\\n- นักเรียน", "detail"),
    ("I2.3", "แสดงยอดเครดิตคงเหลือ", "Summary: ยอดเครดิตทั้งหมด, ใช้ไปแล้ว, คงเหลือ", "detail"),
    ("I2.4", "ติดตามการใช้เครดิต", "แสดง Invoice ที่ใช้ Credit Note หัก\\n- ยอดที่ใช้ / ยอดคงเหลือ", "detail"),

    ("J", "การตั้งค่าบัญชี (Account Settings)", "", "module"),
    ("J1", "จัดการบัญชีผู้ใช้ (Account Management)", "", "sub"),
    ("J1.1", "เปลี่ยนรหัสผ่าน", "แก้ไขรหัสผ่านจากหน้า Settings", "detail"),
    ("J1.2", "แก้ไขข้อมูลส่วนตัว", "อัปเดตชื่อ, Email, เบอร์โทร", "detail"),
    ("J1.3", "ตั้งค่าภาษา", "เลือกภาษาที่แสดงผล: ไทย / อังกฤษ / จีน", "detail"),
    ("J1.4", "ตั้งค่าการแจ้งเตือน", "เปิด/ปิดการแจ้งเตือนตามประเภท", "detail"),

    ("K", "รองรับหลายภาษา (Multi-Language Support)", "", "module"),
    ("K1", "ระบบ i18n", "", "sub"),
    ("K1.1", "รองรับ 2 ภาษา", "ภาษาที่รองรับ:\\n1. อังกฤษ (English)\\n2. ไทย (Thai)", "detail"),
    ("K1.2", "เปลี่ยน Font อัตโนมัติตามภาษา", "Font:\\n- ไทย: Sukhumvit\\n- อังกฤษ: Lato", "detail"),
    ("K1.3", "แสดงสกุลเงินตาม Locale", "รูปแบบ: ฿XX,XXX.XX\\n- ใช้ Locale-aware formatting", "detail"),

    ("L", "Responsive Design", "", "module"),
    ("L1", "Mobile Components", "", "sub"),
    ("L1.1", "Mobile Bottom Navigation", "แถบ Tab navigation ด้านล่างหน้าจอ\\n- ไอคอน + ชื่อ Tab\\n- Highlight Tab ที่ active", "detail"),
    ("L1.2", "Mobile Cart Drawer", "Slide-out sidebar สำหรับตะกร้า\\n- แสดงรายการ, ลบรายการ\\n- ยอดรวม + ปุ่ม Checkout", "detail"),
    ("L1.3", "Mobile Filter Section", "ตัวกรองแบบ compact สำหรับ Mobile\\n- Touch-friendly buttons", "detail"),
    ("L1.4", "Mobile Summary Carousel", "Horizontal scrolling summary boxes\\n- Swipe เพื่อดู metric ต่างๆ", "detail"),
    ("L1.5", "Skeleton Loading", "แสดง Skeleton UI ระหว่างโหลดข้อมูล\\n- Placeholder cards และ filter", "detail"),
]

print(f"Total data rows: {len(rows_data)}")

# ── Create header row + main table ────────────────────────────────

# +1 for header row
total_table_rows = 1 + len(rows_data)
print(f"Total table rows (including header): {total_table_rows}")

table = doc.add_table(rows=total_table_rows, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
WBS_WIDTH = 710
FEATURE_WIDTH = 4150
DESC_WIDTH = 5900

# Header row
header_cells = table.rows[0].cells
for i, header_text in enumerate(["WBS", "Feature", "Description"]):
    set_cell_shading(header_cells[i], "D9D9D9")
    write_cell_text(header_cells[i], header_text, bold=True, size=14)
    set_cell_vertical_alignment(header_cells[i])

# Set column widths on header
header_cells[0].width = Twips(WBS_WIDTH)
header_cells[1].width = Twips(FEATURE_WIDTH)
header_cells[2].width = Twips(DESC_WIDTH)

# Populate data rows
for row_idx, (wbs, feature, desc, row_type) in enumerate(rows_data):
    row = table.rows[row_idx + 1]  # +1 for header
    cells = row.cells

    # Set widths
    cells[0].width = Twips(WBS_WIDTH)
    cells[1].width = Twips(FEATURE_WIDTH)
    cells[2].width = Twips(DESC_WIDTH)

    if row_type == "module":
        # Grey shading on all cells, bold text
        for cell in cells:
            set_cell_shading(cell, "D9D9D9")
        write_cell_text(cells[0], wbs, bold=True, size=14)
        write_cell_text(cells[1], feature, bold=True, size=14)
        write_cell_text(cells[2], desc, bold=True, size=14)
    elif row_type == "sub":
        # Bold WBS and Feature, no shading
        write_cell_text(cells[0], wbs, bold=True, size=14)
        write_cell_text(cells[1], feature, bold=True, size=14)
        write_cell_text(cells[2], desc, bold=False, size=14)
    else:
        # Detail row - normal weight
        write_cell_text(cells[0], wbs, bold=False, size=14)
        write_cell_text(cells[1], feature, bold=False, size=14)
        write_cell_text(cells[2], desc, bold=False, size=14)

    # Set vertical alignment
    for cell in cells:
        set_cell_vertical_alignment(cell)

# ── Set table column widths via tblGrid ───────────────────────────
tbl = table._tbl
tblGrid = tbl.find(qn('w:tblGrid'))
if tblGrid is not None:
    gridCols = tblGrid.findall(qn('w:gridCol'))
    widths = [WBS_WIDTH, FEATURE_WIDTH, DESC_WIDTH]
    for gc, w in zip(gridCols, widths):
        gc.set(qn('w:w'), str(w))

# ── Save document ─────────────────────────────────────────────────

doc.save(output_path)
print(f"\nDocument saved to: {output_path}")
print(f"Total rows in main table (including header): {total_table_rows}")
print(f"Total data rows (excluding header): {len(rows_data)}")

# Count by type
module_count = sum(1 for r in rows_data if r[3] == "module")
sub_count = sum(1 for r in rows_data if r[3] == "sub")
detail_count = sum(1 for r in rows_data if r[3] == "detail")
print(f"\nBreakdown:")
print(f"  Module headers (grey): {module_count}")
print(f"  Sub-level rows (bold): {sub_count}")
print(f"  Detail rows (normal): {detail_count}")
print(f"  Total: {module_count + sub_count + detail_count}")
